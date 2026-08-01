import json
import re
from docx import Document

dictionary = {}

with open("CUADv1.json" , "r", encoding = "utf-8") as file: #einlesen json datei
    summary = json.load(file)
    title_pattern = re.compile(r"[_-]+")
    def experiment():
        first_contract = summary['data'][0] #erster vertrag im json dic
        #for key in first_contract: #keys von datei ausfindig machen
        #    print(key)             #aufgeteilt in 'title' und 'paragraphs'
        text = first_contract.get('title')
        teile1 = re.split(r"_", text)
        for i in range (len(teile1)):
            teile2 = re.split(r"-", teile1[i])
        teile3 = re.split(r"-", text)
        #zahlen = re.findall(r"\d", text)
        #trennen = re.sub()
        teile1.pop(len(teile1) - 1)
        str; Contract_Type1 = teile1[len(teile1) - 1]
        str; Contract_Type2 = teile3[len(teile3) - 1] 
        #print("Ausgabe1: ", Contract_Type1, " Ausgabe2: ", Contract_Type2, "\n")
        #print("1. title: \n", teile1 + teile2)
        #print("Zahlen des Dokuments: \n", zahlen)
        #print("1. title: \n", teile2)
        #print("title: ", first_contract.get('title')) 
        # Target: reverse list, extract whole element with Agreement in it, write it into type
        all_pieces = teile1 + teile2
        all_pieces.reverse()
        #print(all_pieces[0])
    def create_contract_list():
    # # [Company]_[Date]_[SEC Filing]_[Exhibit]_[ContractType]
    #     list; relevant_contracts = []
    #     int; counter = 0
    #     int; service_counter = 0
    #     int; outsourcing_counter = 0
    #     int; supply_counter = 0
    #     int; distributor_counter = 0
    #     int; license_counter = 0
    #     license_indices: list = []
    #     distributor_indices: list = []
    #     service_indices: list = []
    #     supply_indices: list = []
    #     outsourcing_indices: list = []
    #     list; relevant_contract_indices = []
    #     #suche speziell nach: License Agreement, Service Agreement, Supply Agreement, Distributor Agreement and Outsourcing Agreement
    #     for i in range(510):               #Ausgabe aller Titel aller Verträge
    #         list; contract = summary['data'][i]
    #         str; current_contract = contract.get('title')
    #         #list; parts_of_contract = title_pattern.split(current_contract)
    #         #index = parts_of_contract.index("Agreement")
    #         #contract_type = parts_of_contract[index]
    #         #str; contract_type = parts_of_contract[len(parts_of_contract) - 1]
    #         #str; company_of_contract = parts_of_contract[0]
    #         #exact_contract_type = re.findall(r"\w+", contract_type)
    #         str; low = current_contract.lower()
    #         if re.search(r"license|licence|licensing", low) or re.search(r"service(:?s)?|servicing", low) or re.search(r"distribut(?:or|ion)", low) \
    #         or low.find("supply") != -1 or low.find("outsourcing") != -1:
    #             if re.search(r"license|licence|licensing", low):
    #                 int; license_counter += 1
    #                 license_indices.append(i+1)
    #             if re.search(r"service(:?s)?|servicing", low):
    #                 int; service_counter += 1
    #                 service_indices.append(i+1)
    #             if low.find("supply") != -1:
    #                 int; supply_counter += 1
    #                 supply_indices.append(i+1)
    #             if re.search(r"distribut(?:or|ion)", low):
    #                 int; distributor_counter += 1
    #                 distributor_indices.append(i+1)
    #             if low.find("outsourcing") != -1:
    #                 int; outsourcing_counter += 1
    #                 outsourcing_indices.append(i+1)
        #hardcoded lists, if time -> rewrite
        service_indices = [18,25, 49, 54, 56, 59, 76, 79, 104, 123, 161, 180, 201, 211, 213, 236, 249, 279, 286, 310, 326, 357, 386, 407, 456, 458, 461, 507]
        distributor_indices = [1, 68, 86, 89, 90, 102, 121, 135, 137, 154, 156, 165, 168, 199, 217, 241, 261, 263, 266, 298, 311, 312, 341, 346, 360, 367, 370, 380, 448, 459, 472, 494]
        supply_indices = [3, 17, 55, 71, 108, 116, 144, 170, 184, 189, 220, 240, 243, 314, 377, 379, 400, 418]
        outsourcing_indices = [42, 63, 66, 69, 70, 100, 133, 176, 267, 305, 338, 349, 392, 401, 444, 479, 482, 487]
        license_indices = [28, 40, 44, 78, 95, 113, 147, 159, 164, 175, 192, 229, 231, 248, 250, 281, 307, 328, 344, 369, 389, 397, 399, 414, 430, 433, 437, 446, 470, 475, 490, 501, 506]
        #hardcoded counter
        outsourcing_counter = len(outsourcing_indices)
        service_counter = len(service_indices)
        distributor_counter = len(distributor_indices)
        supply_counter = len(supply_indices)
        license_counter = len(license_indices)

        relevant_contract_indices: list = sorted(service_indices + distributor_indices + supply_indices + outsourcing_indices + license_indices)

        dev_of_dcontracts: int = (distributor_counter - 32)
        dev_of_scontracts: int = (service_counter - 28)
        dev_of_supcontracts: int = (supply_counter - 18)
        dev_of_ocontracts: int = (outsourcing_counter - 18)
        dev_of_lcontracts: int = (license_counter - 33)
        print("The relevant contracts are at the following positions: ", relevant_contract_indices)

        counter: int = len(relevant_contract_indices)

        for i in range(len(relevant_contract_indices)):
            list; contract = summary['data'][relevant_contract_indices[i] - 1]
            str; current_contract = contract.get('title')
            print("Title of contract no. %d: " % (relevant_contract_indices[i]),  current_contract, "\n")

        print("You can use %d" % (counter), "different contracts for training", " split up in: \n",\
            "distributor: %d," % distributor_counter, "perfectly fine" if (dev_of_dcontracts == 0) else ("meaning the count is off by %d -> false positives" % dev_of_dcontracts) if dev_of_dcontracts > 0 else ("meaning the count is off by %d -> false negatives" % dev_of_dcontracts) ,"\n",\
            "service: %d," % service_counter, "perfectly fine" if (dev_of_scontracts == 0) else ("meaning the count is off by %d -> false positives" % dev_of_scontracts) if dev_of_scontracts > 0 else ("meaning the count is off by %d -> false negatives" % dev_of_scontracts) ,"\n",\
            "supply: %d," % supply_counter, "perfectly fine" if (dev_of_supcontracts == 0) else ("meaning the count is off by %d -> false positives" % dev_of_supcontracts) if dev_of_supcontracts > 0 else ("meaning the count is off by %d -> false negatives" % dev_of_supcontracts) ,"\n",\
            "outsourcing: %d," % outsourcing_counter, "perfectly fine" if (dev_of_ocontracts == 0) else ("meaning the count is off by %d -> false positives" % dev_of_ocontracts) if dev_of_ocontracts > 0 else ("meaning the count is off by %d -> false negatives" % dev_of_ocontracts) ,"\n",\
            "license: %d," % license_counter, "perfectly fine" if (dev_of_lcontracts == 0) else ("meaning the count is off by %d -> false positives" % dev_of_lcontracts) if dev_of_lcontracts > 0 else ("meaning the count is off by %d -> false negatives" % dev_of_lcontracts) ,"\n")
        print(f"Licenses indices: {license_indices} \n Supply indices: {supply_indices} \n Distributor indices: {distributor_indices} \n Outsourcing indices: {outsourcing_indices} \n Service indices: {service_indices}")
        all_contract_numbers_and_titles(relevant_contract_indices)
    
    def generate_word():  
        #aufbereiten von text
        for i in range(10):
            dict; word_output = summary['data'][i]
            str; word_title = word_output.get('title')
            list; word_title2 = title_pattern.split(word_title)
            str; contract_type_word = word_title2[len(word_title2) - 1]
            str; company_of_contract_word = word_title2[0]
            str; word_paragraphs = word_output.get('paragraphs')
            for item in word_paragraphs:
                str; word_context = item['context']
            #word datei erstellen
            doc = Document()

            doc.add_page_break

            doc.add_heading(f"{contract_type_word} of {company_of_contract_word}", level = 1)

            doc.add_paragraph(word_context)

            doc.save(fr"C:\Users\finnd\Documents\Informatik\Projektarbeit\Contracts\Contract no. %d {contract_type_word} + {company_of_contract_word}.docx" % (i+1))
    #generate_word()
    # function to in general extract content + title and major information of the document
    def gain_access(i: int) -> tuple[str, str, str]:
        output: dict = summary['data'][i]
        title: str = output.get('title')
        title_parts: list = title_pattern.split(title)
        contract_type: str = title_parts[len(title_parts) - 1]
        parentcompany_of_contract: str = title_parts[0]
        paragraphs: list = (output.get('paragraphs'))
        for item in paragraphs:
            str; content = item['context']
        # for i in range()
        # new_content = content.replace("Page -1-", "").replace("Page -2-", "")
        return contract_type, parentcompany_of_contract, content

    def generate_specific_word(text: str, contract_no: int):  
        dict; word_output = summary['data'][contract_no]
        str; word_title = word_output.get('title')
        list; word_title2 = title_pattern.split(word_title)
        str; contract_type_word = word_title2[len(word_title2) - 1]
    
        doc = Document()

        doc.add_page_break

        doc.add_heading(f"Terms for {contract_type_word} type ahh contracts", level = 1)

        doc.add_paragraph(text)

        doc.save(fr"C:\Users\finnd\Documents\Informatik\Projektarbeit\LegalRiskLexicon\Contract Terms for {contract_type_word} type ahh contracts.docx")

    def get_lexiconterms():
        text: str = str(summary['data'][0])
        generate_specific_word(text, 0)
        #print(text)

#    get_lexiconterms()

    def word_of_contract_listing(text: str):  
    
        doc = Document()

        doc.add_page_break()

        doc.add_heading(f"All Contracts and their index put together", level = 1)

        doc.add_paragraph(text)

        doc.save(fr"C:\Users\finnd\Documents\Wirtschaftsinformatik\Informatik\Projektarbeit\Listing_contracts\Listing_of_all_contracts_plus_index_frthistime.docx")

    def all_contract_numbers_and_titles(relevant_contracts: list) -> str:
        text_liste: list = []
        for i in range(len(relevant_contracts)):               #Ausgabe aller Titel aller Verträge
            list; contract = summary['data'][relevant_contracts[i] - 1]
            str; current_contract = contract.get('title')
            text_liste.append(f"Title of contract no. {relevant_contracts[i]} {current_contract}")
        fertige_ausgabe: str = "\n\n".join(text_liste)
        word_of_contract_listing(fertige_ausgabe)

create_contract_list()
#all_contract_numbers_and_titles()


    # create_contract_list()
    # generate_word()

    #|- summary (dict)
#       |- data" (list with dictionaries as entries)
#               |- title (string)
#               |- paragraphs (list)
#                   |- context (string)
#               |- qad (dict)



        #else:
        #    print("Contract no. %d: " % (i+1), " doesn't have the right type")
    # contract = summary['data'][488]
    # current_contract = contract.get('title')
    # parts_of_contract = title_pattern.split(current_contract)
    # print("Whole_contract: ", parts_of_contract)
        #print(first_part, "\n")
        # for y in range(len(first_part) - 1):
        #     second_part = re.split(r"-", first_part[y])
        #print(second_part, "\n")
        # full_current_contract = first_part
        # for x in range(len(full_current_contract) - 1):
        #     final_part = re.split(r"-", full_current_contract[x])
        #print(full_current_contract, "\n")
        # final_part.reverse()
        # full_current_contract.reverse()
    #print(first_contract.get('paragraphs'))
    #for key in first_contract.get('paragraphs'):
    #    print(key)
    #c_paragraphs = first_contract.get('paragraphs')
    #print(c_paragraphs.count('text'))
    #for item in c_paragraphs:
       # print("Content: ", item['context'])
    
